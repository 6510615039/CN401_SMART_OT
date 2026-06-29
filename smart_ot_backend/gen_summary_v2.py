# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'TH Sarabun New'
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(4)

def add_heading_colored(doc, text, color=0xB8001F, size=18):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor((color>>16)&0xFF, (color>>8)&0xFF, color&0xFF)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'B8001F'})
        shading.append(bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(13)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.bold = True
            # สีเขียวสำหรับคอลัมน์สถานะ
            if headers[-1] == 'สถานะ' and ci == len(headers)-1:
                if 'เสร็จ' in str(val) or 'ใช้ได้' in str(val):
                    r.font.color.rgb = RGBColor(0x0A, 0x8A, 0x44)
                    r.bold = True
                elif 'ยังไม่' in str(val):
                    r.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
                    r.bold = True
    for i, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)
    return table

def para(doc, text, bold=False, size=14, color=None, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    if bold: r.bold = True
    if color: r.font.color.rgb = RGBColor((color>>16)&0xFF, (color>>8)&0xFF, color&0xFF)
    if align: p.alignment = align

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

# ═══════════════════════════════════════════════════════════════
# TITLE
# ═══════════════════════════════════════════════════════════════
para(doc, 'สรุปการแก้ไขระบบ SMART OT', bold=True, size=22, color=0xB8001F, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, 'ตามมติการประชุม — สถานะล่าสุด', bold=True, size=16, color=0x333333, align=WD_ALIGN_PARAGRAPH.CENTER)
para(doc, 'วันที่ 27 มิถุนายน 2569', size=14, color=0x666666, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
# SECTION 0: ซ่อมไฟล์
# ═══════════════════════════════════════════════════════════════
add_heading_colored(doc, '0. ซ่อมไฟล์ที่เสียหาย (Merge Conflict + Truncated)')
para(doc, 'แก้ไขไฟล์ที่มี conflict marker ค้าง หรือถูกตัดท้ายจาก merge conflict เดิม รวม 10+ ไฟล์')
para(doc, 'Backend:', bold=True)
for f in ['models.py — Notification model', 'serializers.py — NotificationSerializer',
          'urls.py — routes notifications/checker/head', 'views.py — Notification API + head_report_pdf',
          'settings.py — ส่วนท้ายที่ถูกตัด']:
    bullet(doc, f)
para(doc, 'Frontend:', bold=True)
for f in ['AppShell.tsx — sidebar + dropdown', 'Login.tsx — ForgotPassword component',
          'shared.tsx — SectionCard component', 'executive.tsx — BarChart component',
          'OTDetailPage.tsx — ลบ null bytes', 'deptrep.tsx — RepHistory component']:
    bullet(doc, f)

# ═══════════════════════════════════════════════════════════════
# SECTION 1: ADMIN
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_colored(doc, '1. แอดมิน (Admin)')

add_table(doc,
    ['#', 'รายการ', 'รายละเอียด', 'สถานะ'],
    [
        ['A1', 'ป้องกันอัปโหลดไฟล์เวลาซ้ำ', 'บล็อกทันทีถ้าเดือนนั้นมีข้อมูลแล้ว ลบ dialog ยืนยันนำเข้าซ้ำออกทั้งหมด', 'เสร็จแล้ว'],
        ['A2', 'แก้ไขข้อมูลเวลารายบุคคล', 'AdminEditableTable — มีอยู่แล้ว', 'เสร็จแล้ว'],
        ['A3', 'Hardcode เพดาน OT + ตัดเศษนาที', 'วันธรรมดา 4 ชม. วันหยุด 7 ชม. ใช้เฉพาะชั่วโมงเต็ม หน้าตั้งค่าแสดงแบบอ่านอย่างเดียว', 'เสร็จแล้ว'],
        ['A4', 'บันทึก rate_per_hour ลงคำร้อง', 'ป้องกันผลกระทบ Historical Data เมื่ออัตราเปลี่ยน', 'เสร็จแล้ว'],
        ['A5', 'ลบวันหยุดราชการ (is_system) ได้', 'แอดมินลบได้ role อื่นลบไม่ได้ เพิ่มปุ่มลบในหน้าจัดการวันหยุด', 'เสร็จแล้ว'],
        ['A6', 'ตั้ง Cut-off Date', 'OTDeadline — มีอยู่แล้ว', 'เสร็จแล้ว'],
        ['A7', 'เพิ่มคอลัมน์ "กะ" ในตาราง', 'แสดงกะเช้า/กะปกติ + backend ส่ง timePeriod กลับมา', 'เสร็จแล้ว'],
        ['A8', 'วันที่แสดงเป็น DD-MM-YY', 'เพิ่ม formatDateDDMMYY() ในตารางข้อมูลเวลา', 'เสร็จแล้ว'],
        ['A9', 'Import รายชื่อพนักงานจาก Excel', 'management command import_staff_master อ่าน 2 ไฟล์ สร้าง User+แผนก+อีเมล', 'เสร็จแล้ว'],
        ['A10', 'จำกัด login เฉพาะคนที่มีบัญชี', 'ไม่สร้าง user อัตโนมัติจาก TU API อีก ต้องให้ admin สร้างก่อน', 'เสร็จแล้ว'],
        ['A11', 'ลบแผนกที่ไม่อยู่ใน Excel', 'เหลือ 6 แผนกตามโครงสร้างจริง', 'เสร็จแล้ว'],
        ['A12', 'แสดงผู้ใช้ทั้งหมด (ไม่ pagination)', 'ปิด pagination ใน UserViewSet', 'เสร็จแล้ว'],
    ],
    [1.0, 4.0, 6.5, 2.5]
)

# ═══════════════════════════════════════════════════════════════
# SECTION 2: CHECKER
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_colored(doc, '2. ผู้ตรวจสอบ (Checker)')

add_table(doc,
    ['#', 'รายการ', 'รายละเอียด', 'สถานะ'],
    [
        ['C1', 'ไม่ถูกจำกัดด้วย deadline', 'ตรวจสอบแล้วใช้ได้อยู่แล้ว (deadline อยู่ในขั้นตอนยื่นคำร้อง)', 'ใช้ได้อยู่แล้ว'],
        ['C2', 'แสดงแผนกที่ไม่มีการเบิก OT', 'เพิ่มกรอบแจ้งเตือนสีแดงใน Dashboard + เรียก API no-ot-departments', 'เสร็จแล้ว'],
        ['C3', 'สีสถานะ "ยังไม่ส่งเอกสาร"', 'ใช้สีส้ม (warning) ไม่ซ้ำกับ reject (แดง)', 'เสร็จแล้ว'],
        ['C4', 'ปรับ UI ปุ่มอนุมัติ/ตีกลับ', 'ขยายปุ่ม เพิ่มข้อความกำกับจำนวนรายการ', 'เสร็จแล้ว'],
        ['C5', 'เพิ่ม % บนกราฟแท่ง', 'แสดงเปอร์เซ็นต์สัดส่วนแต่ละแผนกบนหัวแท่ง', 'เสร็จแล้ว'],
    ],
    [1.0, 4.5, 6.5, 2.5]
)

# ═══════════════════════════════════════════════════════════════
# SECTION 3: DEPTREP
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_colored(doc, '3. ตัวแทนฝ่าย (DeptRep)')

add_table(doc,
    ['#', 'รายการ', 'รายละเอียด', 'สถานะ'],
    [
        ['D1', 'ปุ่มอัปโหลดเอกสารใน Step 3', 'รองรับ .xlsx/.pdf สำหรับกรณีจัดฟอร์แมตเอง', 'เสร็จแล้ว'],
        ['D2', 'เพิ่มแถว "ยอดยกมา"', 'แสดงทั้ง Preview และ Excel', 'เสร็จแล้ว'],
        ['D3', 'เปลี่ยนชื่อผู้ลงนาม', 'เพิ่มช่องกรอกชื่อผู้ลงนามใน Step 3', 'เสร็จแล้ว'],
        ['D4', 'Excel ตรงแบบฟอร์มจริง', '16 คอลัมน์ (A-P) / 8 คอลัมน์วันที่ / ช่องลงนาม 2 ฝั่ง / Sheet "หลักฐานจ่าย"', 'เสร็จแล้ว'],
        ['D5', 'ลบ filter เดือน ออก', 'แสดงคำร้องรออนุมัติทั้งหมดโดยไม่ต้อง filter', 'เสร็จแล้ว'],
        ['D6', 'ล็อกอีเมลผู้ตรวจสอบ', 'ใช้ @reg.tu.ac.th (notify_email) เท่านั้น', 'เสร็จแล้ว'],
    ],
    [1.0, 4.5, 6.5, 2.5]
)

# ═══════════════════════════════════════════════════════════════
# SECTION 4: ข้าม ROLE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_colored(doc, '4. ข้าม Role (ทั้งระบบ)')

add_table(doc,
    ['#', 'รายการ', 'รายละเอียด', 'สถานะ'],
    [
        ['X1', 'Notification อัตโนมัติ', 'สร้าง in-app notification ทุกครั้งที่สถานะเปลี่ยน (ยื่น/อนุมัติ/ตีกลับ/ส่งต่อ)', 'เสร็จแล้ว'],
        ['X2', 'Email แจ้งเตือน', 'ส่ง email ทุกครั้งที่สถานะเปลี่ยน รองรับ SMTP จริง (Gmail)', 'เสร็จแล้ว (รอใส่ SMTP)'],
        ['X3', 'Login จำกัดเฉพาะพนักงาน', 'TU API ยืนยันตัวตนได้ แต่ต้องมีบัญชีในระบบก่อน', 'เสร็จแล้ว'],
        ['X4', 'Login ลืมรหัสผ่าน', 'ลิงก์ไป TU IT Center (account.tu.ac.th) โดยตรง', 'เสร็จแล้ว'],
        ['X5', 'Role terminology มาตรฐาน', 'พนักงาน/หัวหน้างาน/ตัวแทนฝ่าย/ผู้ตรวจสอบ/ผู้บริหาร/แอดมิน', 'เสร็จแล้ว'],
        ['X6', 'Breadcrumbs หัวหน้างาน', 'มีอยู่แล้ว', 'เสร็จแล้ว'],
        ['X7', 'ปุ่มอนุมัติหัวหน้างาน', 'เปลี่ยนจากเครื่องหมาย ✓ เป็นคำว่า "อนุมัติ"', 'เสร็จแล้ว'],
        ['X8', 'แยก filter เดือน/ปี', 'หน้า Staff timelog แยก dropdown เดือน + ปี', 'เสร็จแล้ว'],
        ['X9', 'ตั้งค่า OT อ่านอย่างเดียว', 'หน้าตั้งค่าระบบแสดงเพดาน/อัตราเป็นข้อความ ไม่ให้แก้', 'เสร็จแล้ว'],
    ],
    [1.0, 4.5, 6.5, 2.5]
)

# ═══════════════════════════════════════════════════════════════
# SECTION 5: ยังไม่ได้ทำ
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_colored(doc, '5. สิ่งที่ยังไม่ได้ทำ / รอดำเนินการ', color=0xD32F2F)

add_table(doc,
    ['#', 'รายการ', 'รายละเอียด', 'หมายเหตุ'],
    [
        ['P1', 'ใส่ SMTP จริง (.env)', 'ต้องสร้าง Gmail App Password แล้วใส่ใน .env', 'รอ credentials'],
        ['P2', 'Docker + Deploy', 'ศึกษา Docker build + ประสานคุณบอลเรื่อง Server/VPN', 'ตามมติประชุม'],
        ['P3', 'ปรับ Frontend เป็น Django template', 'ลด React พึ่งพา เพื่อส่งมอบง่าย', 'ตามมติประชุม'],
        ['P4', 'แสดงตำแหน่งการจ้างงานในโปรไฟล์', 'แสดงตำแหน่งจริง ไม่ใช่แค่ชื่อแผนก', 'ยังไม่ได้ทำ'],
        ['P5', 'อัปโหลดรูปโปรไฟล์', 'ให้ผู้ใช้อัปโหลดรูปภาพโปรไฟล์ตัวเอง', 'ยังไม่ได้ทำ'],
        ['P6', 'วันที่ DD-MM-YY ทั้งระบบ', 'ตอนนี้แก้เฉพาะตาราง import ยังไม่ครอบคลุมทุกหน้า', 'ทำบางส่วน'],
        ['P7', 'Excel ตัดแบ่ง >8 วัน', 'กรณีพนักงานยื่น OT เกิน 8 วัน/เดือน ต้องแบ่งแถว', 'ยังไม่ได้ทำ'],
        ['P8', 'ทดสอบ Notification ครบ flow', 'ยื่น -> อนุมัติ -> ส่งต่อ -> ตรวจสอบ ทุกขั้นตอน', 'รอทดสอบ'],
        ['P9', 'Mapping email @tu.ac.th กับ @reg', 'จับคู่อีเมลสำหรับรับแจ้งเตือน', 'เสร็จแล้วใน import'],
    ],
    [1.0, 4.5, 6.0, 3.0]
)

# ═══════════════════════════════════════════════════════════════
# SECTION 6: สถานะ Build
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_colored(doc, '6. สถานะการตรวจสอบ')

add_table(doc,
    ['รายการ', 'ผลลัพธ์', 'สถานะ'],
    [
        ['Vite Build (Frontend)', 'Build สำเร็จ ไม่มี error', 'ผ่าน'],
        ['Python Syntax (Backend)', 'Compile สำเร็จทุกไฟล์', 'ผ่าน'],
        ['Django Migrate', 'Migration สำเร็จ (Notification model)', 'ผ่าน'],
        ['Import พนักงาน 53 คน', '52 สร้างใหม่ + merge แผนก/อีเมลครบ', 'ผ่าน'],
    ],
    [5.0, 5.5, 3.5]
)

# ═══════════════════════════════════════════════════════════════
# FOOTER
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
para(doc, 'รวมทั้งสิ้น: เสร็จแล้ว 28 รายการ / รอดำเนินการ 9 รายการ',
     bold=True, size=16, color=0xB8001F, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para(doc, 'จัดทำโดย: ทีมพัฒนาระบบ SMART OT', size=13, color=0x666666, align=WD_ALIGN_PARAGRAPH.RIGHT)
para(doc, 'วันที่ 27 มิถุนายน 2569', size=13, color=0x666666, align=WD_ALIGN_PARAGRAPH.RIGHT)

output = r'C:\Users\ASUS\Desktop\project\สรุปการแก้ไข_SMART_OT_v2.docx'
doc.save(output)
print(f'Created: {output}')
