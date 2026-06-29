from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
style.font.name = 'TH Sarabun New'
style.font.size = Pt(14)
style.paragraph_format.space_after = Pt(4)

# -- Title --
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('สรุปการแก้ไขระบบ SMART OT\nตามมติการประชุม')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0xB8, 0x00, 0x1F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('วันที่ 27 มิถุนายน 2569')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacer

# -- Helper: add table --
def add_table(doc, headers, rows, col_widths_cm):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # bg color
        shading = cell._element.get_or_add_tcPr()
        bg = shading.makeelement(qn('w:shd'), {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'B8001F',
        })
        shading.append(bg)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(14)
            if ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r.bold = True

    # Column widths
    for i, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    return table

# -- Helper: section heading --
def add_heading_colored(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0xB8, 0x00, 0x1F)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)

# ═══════════════════════════════════════════════════════════════
# Section 1: Fix files
# ═══════════════════════════════════════════════════════════════
add_heading_colored(doc, '1. ซ่อมไฟล์ที่เสียหาย (Merge Conflict + Truncated Files)')

doc.add_paragraph(
    'แก้ไขไฟล์ที่มี conflict marker ค้าง หรือถูกตัดท้ายจาก merge conflict เดิม รวมทั้งหมด 10 ไฟล์'
)

p = doc.add_paragraph()
r = p.add_run('Backend (4 ไฟล์):')
r.bold = True
r.font.size = Pt(14)

for item in [
    'models.py — เพิ่ม Notification model ที่ถูกตัด',
    'serializers.py — เพิ่ม NotificationSerializer ที่ถูกตัด',
    'urls.py — เพิ่ม routes สำหรับ notifications, checker budget, head report',
    'views.py — เพิ่ม Notification API, head_report_pdf_view ที่ถูกตัด',
]:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
r = p.add_run('Frontend (6 ไฟล์):')
r.bold = True
r.font.size = Pt(14)

for item in [
    'AppShell.tsx — เขียน sidebar, content area, dropdown ที่ถูกตัด',
    'Login.tsx — เขียน ForgotPassword component ที่ถูกตัด',
    'shared.tsx — เขียน SectionCard component ที่ถูกตัด',
    'executive.tsx — เขียน BarChart component ที่ถูกตัด',
    'OTDetailPage.tsx — ลบ null bytes ที่ต่อท้ายไฟล์',
    'deptrep.tsx — เขียน RepHistory component ที่ถูกตัด',
]:
    doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════════════════════════════════════
# Section 2: Admin
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_colored(doc, '2. แอดมิน (Admin) — 4 รายการ')
doc.add_paragraph('การแก้ไขเกี่ยวกับสิทธิ์และฟังก์ชันของแอดมิน ตามมติการประชุม')

add_table(doc,
    ['รหัส', 'รายการ', 'รายละเอียด', 'ไฟล์ที่แก้'],
    [
        ['A1', 'ป้องกันอัปโหลดไฟล์เวลาซ้ำเดือนเดียวกัน',
         'ระบบตรวจสอบว่าเดือนที่นำเข้ามีข้อมูลอยู่แล้วหรือไม่ หากซ้ำจะแจ้งเตือนให้ยืนยัน (HTTP 409)', 'views.py'],
        ['A3', 'Hardcode เพดาน OT + ตัดเศษนาที',
         'วันธรรมดาสูงสุด 4 ชม. วันหยุดสูงสุด 7 ชม. ใช้เฉพาะชั่วโมงเต็ม', 'views.py'],
        ['A4', 'บันทึกอัตราค่าตอบแทน (rate_per_hour)',
         'เก็บอัตราลงในคำร้อง เพื่อป้องกันผลกระทบต่อข้อมูลย้อนหลังเมื่ออัตราเปลี่ยน', 'views.py\nserializers.py'],
        ['A5', 'อนุญาตแอดมินลบวันหยุดราชการ (is_system)',
         'แอดมินลบวันหยุดจากปฏิทินระบบได้ role อื่นยังลบไม่ได้', 'views.py'],
    ],
    [1.5, 4.5, 6.5, 3.0]
)

# ═══════════════════════════════════════════════════════════════
# Section 3: Checker
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_colored(doc, '3. ผู้ตรวจสอบ (Checker) — 5 รายการ')
doc.add_paragraph('การแก้ไขเกี่ยวกับสิทธิ์และหน้าจอของผู้ตรวจสอบ ตามมติการประชุม')

add_table(doc,
    ['รหัส', 'รายการ', 'รายละเอียด', 'สถานะ'],
    [
        ['C1', 'Checker ไม่ถูกจำกัดด้วย deadline',
         'ตรวจสอบแล้วใช้ได้อยู่แล้ว (deadline check อยู่ในขั้นตอนยื่นคำร้องเท่านั้น)', 'ใช้ได้อยู่แล้ว'],
        ['C2', 'แสดงแผนกที่ไม่มีการเบิก OT',
         'เพิ่มกรอบแจ้งเตือนสีแดงใน Dashboard พร้อมรายชื่อแผนก', 'เสร็จแล้ว'],
        ['C3', 'ปรับสีสถานะ "ยังไม่ส่งเอกสาร"',
         'ใช้สีส้ม (warning) แทนสีเทา ไม่ซ้ำกับสถานะตีกลับ (reject)', 'เสร็จแล้ว'],
        ['C4', 'ปรับปรุง UI ปุ่มอนุมัติ/ตีกลับ',
         'ขยายขนาดปุ่ม เพิ่มข้อความกำกับจำนวนรายการ เพิ่ม shadow', 'เสร็จแล้ว'],
        ['C5', 'เพิ่ม % กำกับบนกราฟแท่ง',
         'แสดงเปอร์เซ็นต์สัดส่วนของแต่ละแผนกบนหัวแท่งกราฟ', 'เสร็จแล้ว'],
    ],
    [1.5, 4.5, 6.5, 3.0]
)

# ═══════════════════════════════════════════════════════════════
# Section 4: DeptRep
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading_colored(doc, '4. ตัวแทนฝ่าย (DeptRep) — 4 รายการ')
doc.add_paragraph('การแก้ไขเกี่ยวกับฟังก์ชันส่งออก Excel และขั้นตอนการส่งต่อเอกสาร ตามมติการประชุม')

add_table(doc,
    ['รหัส', 'รายการ', 'รายละเอียด', 'ไฟล์ที่แก้'],
    [
        ['D1', 'เพิ่มปุ่มอัปโหลดเอกสารใน Step 3',
         'รองรับ .xlsx และ .pdf สำหรับกรณีตัวแทนฝ่ายปรับฟอร์แมตเอง', 'deptrep.tsx'],
        ['D4', 'เพิ่มแถว "ยอดยกมา"',
         'แสดงทั้งในหน้า Preview และไฟล์ Excel export', 'deptrep.tsx'],
        ['D5', 'เพิ่มช่องเปลี่ยนชื่อผู้ลงนาม',
         'ตัวแทนฝ่ายระบุชื่อผู้ลงนามในเอกสาร Excel ได้', 'deptrep.tsx'],
        ['D6', 'ตัดแบ่งข้อมูลเกิน 5 วัน',
         'แบ่งเป็นแถวย่อย (5 วัน/แถว) ตามมาตรฐานงานสารบรรณ', 'deptrep.tsx'],
    ],
    [1.5, 4.5, 6.5, 3.0]
)

# ═══════════════════════════════════════════════════════════════
# Section 5: Verification
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
add_heading_colored(doc, '5. สถานะการตรวจสอบ')

add_table(doc,
    ['รายการตรวจสอบ', 'ผลลัพธ์', 'สถานะ'],
    [
        ['Vite Build (Frontend)', 'Build สำเร็จ ไม่มี error', 'ผ่าน ✓'],
        ['Python Syntax (Backend)', 'Compile สำเร็จทุกไฟล์', 'ผ่าน ✓'],
    ],
    [5.5, 5.5, 3.5]
)

# -- Summary footer --
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('รวมทั้งสิ้น 13 รายการที่ดำเนินการแล้ว')
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0xB8, 0x00, 0x1F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ซ่อมไฟล์ 10 ไฟล์ | Admin 4 รายการ | Checker 5 รายการ | DeptRep 4 รายการ')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('จัดทำโดย: ทีมพัฒนาระบบ SMART OT\nวันที่ 27 มิถุนายน 2569')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.italic = True

# Save
output = r'C:\Users\ASUS\Desktop\project\สรุปการแก้ไข_SMART_OT.docx'
doc.save(output)
print(f'Created: {output}')
